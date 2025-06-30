import os
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from tkinter.scrolledtext import ScrolledText
import subprocess
import platform
import re
from PIL import Image, ImageTk
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from tkinter import filedialog


COLORS = {
    'primary':   '#1e3a8a',
    'secondary': '#3b82f6',
    'accent':    '#10b981',
    'danger':    '#ef4444',
    'warning':   '#f59e0b',
    'background':'#f8fafc',
    'surface':   '#ffffff',
    'text':      '#1f2937',
    'border':    '#e5e7eb'
}


# Add these new imports for summarization
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
import warnings
warnings.filterwarnings("ignore")

# Initialize summarization model (runs locally)
print("Loading summarization model...")
try:
    # Using a lightweight BART model for summarization
    model_name = "facebook/bart-large-cnn"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    summarizer = pipeline("summarization", model=model, tokenizer=tokenizer, 
                         max_length=300, min_length=50, do_sample=False)
    print("Summarization model loaded successfully!")
except Exception as e:
    print(f"Error loading summarization model: {e}")
    summarizer = None

# Folder where all the Word docs are stored
DOCS_FOLDER = "data_folder"

# Initialize root window
root = tk.Tk()
root.title("Well Report Viewer")
root.state('zoomed')
root.geometry("1200x700")


# --- Heading with Image and Title ---
heading_frame = tk.Frame(root, bg="#800000", height=80)
heading_frame.pack(fill="x")
heading_frame.pack_propagate(False)


# Load and resize logo
logo_path = r"C:\Users\bhowm\OneDrive\Desktop\New folder\ongc_logo.png"
logo_image = Image.open(logo_path)
logo_image = logo_image.resize((60, 60), Image.LANCZOS)
logo_photo = ImageTk.PhotoImage(logo_image)

# Place logo on the left
logo_label = tk.Label(heading_frame, image=logo_photo, bg="#800000")
logo_label.image = logo_photo  # Keep reference to avoid garbage collection
logo_label.pack(side="left", padx=10, pady=10)

# Place heading text next to it
heading_label = tk.Label(
    heading_frame,
    text="ONGC Well Perforation & Testing Analysis - Tripura Asset",
    font=("Segoe UI", 18, "bold"),
    bg="#800000",
    fg="white",
    padx=10
)
heading_label.pack(side="left", pady=20)


def set_theme(bg_color, fg_color):
    root.config(bg=bg_color)
    for widget in root.winfo_children():
        try:
            widget.config(bg=bg_color, fg=fg_color)
        except:
            try:
                widget.config(bg=bg_color)
            except:
                pass

# Theme configuration
def toggle_theme():
    if root['bg'] == 'white' or root['bg'] == '#ffffff':
        set_theme('black', 'white')
        theme_btn.config(text='Light Mode', bg='gray', fg='white')
    else:
        set_theme('white', 'black')
        theme_btn.config(text='Dark Mode', bg='lightgray', fg='black')

root.config(bg='white')

theme_btn = tk.Button(root, text="Dark Mode", command=toggle_theme, bg='lightgray', fg='black')
theme_btn.pack(pady=5, anchor='ne')

current_document = None
search_results = []
current_search_index = 0
document_paragraphs = []
  
def export_summary_to_pdf():
    """
    Export the entire summary & table text shown in the GUI to a PDF.
    Uses reportlab.canvas but wraps lines so nothing is cut off.
    """
    if not current_document:
        messagebox.showwarning("No Document",
                               "Please load and analyse a report first.")
        return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf")],
        title="Save Report As PDF"
    )
    if not file_path:
        return

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(file_path, pagesize=letter)
        page_w, page_h = letter
        margin_x = 40
        line_height = 14          # 10-pt font, 4-pt leading
        y = page_h - margin_x

        def wrap_line(text, font_name, font_size, max_width):
            """
            Split 'text' into a list of shorter strings that all fit
            within 'max_width' (in points) for the given font.
            """
            words = text.split(" ")
            lines = []
            cur = ""
            for w in words:
                test = f"{cur} {w}".strip()
                if c.stringWidth(test, font_name, font_size) <= max_width:
                    cur = test
                else:
                    if cur:
                        lines.append(cur)
                    cur = w
            if cur:
                lines.append(cur)
            return lines

        field = field_var.get()
        well = well_var.get()
        keyword = keyword_var.get()

        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin_x, y, "Well Analysis Report")
        y -= line_height * 1.5

        c.setFont("Helvetica", 10)
        meta = f"Field: {field}  |  Well: {well}  |  Keyword: {keyword}"
        for line in wrap_line(meta, "Helvetica", 10, page_w - 2 * margin_x):
            c.drawString(margin_x, y, line)
            y -= line_height
        y -= line_height  

        max_text_width = page_w - 2 * margin_x
        c.setFont("Helvetica", 10)

        for widget, text in perform_smart_search.content_widgets:
            for raw in text.strip().split("\n"):
                if not raw.strip():        
                    y -= line_height
                    continue

                for wrapped in wrap_line(raw.strip(), "Helvetica", 10,
                                         max_text_width):
                    if y < margin_x:        
                        c.showPage()
                        c.setFont("Helvetica", 10)
                        y = page_h - margin_x
                    c.drawString(margin_x, y, wrapped)
                    y -= line_height

            y -= line_height

        c.save()
        messagebox.showinfo("Success",
                            f"Report saved as PDF:\n{file_path}")

    except Exception as e:
        messagebox.showerror("Error", f"Failed to save PDF.\n{e}")

# Populate field and well lists
def get_field_well_map():
    mapping = {}
    
    if not os.path.exists(DOCS_FOLDER):
        print(f"Warning: {DOCS_FOLDER} directory not found.")
        return mapping
    
    print(f"Looking in: {DOCS_FOLDER}")
    print(f"Contents: {os.listdir(DOCS_FOLDER) if os.path.exists(DOCS_FOLDER) else 'Directory not found'}")
    
    for field in os.listdir(DOCS_FOLDER):
        field_path = os.path.join(DOCS_FOLDER, field)
        if not os.path.isdir(field_path):
            continue

        docs = []
        for fname in os.listdir(field_path):
            if not fname.lower().endswith(".docx") or fname.startswith("~$"):
                continue
            full_path = os.path.join(field_path, fname)
            
            try:
                doc = Document(full_path)
                
                well = None
                for p in doc.paragraphs:
                    t = p.text.strip()
                    m = re.match(r"(?i)^Well\s*Name\s*:\s*(.+)", t) or \
                        re.match(r"(?i)^Well\s*:\s*(.+)", t)
                    if m:
                        well = m.group(1).strip()
                        break
                if not well:
                    well = os.path.splitext(fname)[0]

                docs.append((well, full_path))
                print(f"Found well: {well} in field: {field}")
            except Exception as e:
                print(f"Error reading {full_path}: {e}")
                continue

        if docs:
            mapping[field] = docs
    
    print(f"Final mapping: {mapping}")
    return mapping

# Get field-well mapping
field_well_map = get_field_well_map()
fields = list(field_well_map.keys())

# SIDEBAR 
sidebar = tk.Frame(root, bg=COLORS['surface'], relief=tk.SOLID, bd=1)
sidebar.pack(side="left", fill="y", padx=(0,20))
sidebar.config(width=340)

# header
sidebar_header = tk.Frame(sidebar, bg='#820314')
sidebar_header.pack(fill="x")
tk.Label(sidebar_header,
         text="Well Report Control Panel",
         bg='#820314', fg="white",
         font=("Segoe UI", 16, "bold")
).pack(pady=15)

# scrollable area
scrollbar = tk.Scrollbar(sidebar, orient=tk.VERTICAL)
scroll_canvas = tk.Canvas(sidebar,
                          yscrollcommand=scrollbar.set,
                          bg=COLORS['surface'],
                          highlightthickness=0)
scroll_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
scrollbar.config(command=scroll_canvas.yview)

scrollable_window = tk.Frame(scroll_canvas, bg=COLORS['surface'])
scrollable_window.bind(
  "<Configure>",
  lambda e: scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all"))
)
scroll_canvas.create_window((0,0), window=scrollable_window, anchor="nw", width=380)

# bind mousewheel
scroll_canvas.bind_all("<MouseWheel>",
  lambda e: scroll_canvas.yview_scroll(int(-1*(e.delta/120)), "units")
)

# now put your controls into this:
controls_frame = scrollable_window
controls_frame.config(highlightbackground="#800000", highlightthickness=2)

# Field
field_label = tk.Label(controls_frame,
                       text="Field:",
                       bg="white",
                       fg="black",
                       font=("Segoe UI", 10, "bold"))
field_label.pack(anchor="w", pady=(0,2))

field_var = tk.StringVar()
field_dropdown = ttk.Combobox(controls_frame,
                              textvariable=field_var,
                              values=fields,
                              state="readonly",
                              width=25)              # ← 25 chars wide
field_dropdown.pack(fill="x", padx=10, pady=(0,10))

# Well Name
well_label = tk.Label(controls_frame,
                      text="Well Name:",
                      font=("Segoe UI", 10, "bold"))
well_label.pack(anchor="w", pady=(0,2))

well_var = tk.StringVar()
well_dropdown = ttk.Combobox(controls_frame,
                             textvariable=well_var,
                             values=[],
                             state="readonly",
                             width=25)               # ← 25 chars wide
well_dropdown.pack(fill="x", padx=10, pady=(0,10))

# Specific Term
keyword_label = tk.Label(controls_frame,
                         text="Specific Term:",
                         font=("Segoe UI", 10, "bold"))
keyword_label.pack(anchor="w", pady=(0,2))

keyword_var = tk.StringVar()
keyword_dropdown = ttk.Combobox(controls_frame,
                                textvariable=keyword_var,
                                values=["Perforation"],
                                state="readonly",
                                width=25)    
keyword_dropdown.current(0)
keyword_dropdown.pack(fill="x", padx=10, pady=(0,15))

# Search Section
search_label = tk.Label(controls_frame, text="✨ Smart Search & Summary:", font=("Segoe UI", 10, "bold"))
search_label.pack(anchor="w")

search_var = tk.StringVar()
search_entry = tk.Entry(controls_frame, textvariable=search_var, font=("Segoe UI", 10))
search_entry.pack(fill="x", padx=10, pady=5)

search_btn_frame = tk.Frame(controls_frame)
search_btn_frame.pack(fill="x", padx=10, pady=2)


# --- UI Layout (RIGHT) ---
content_frame = tk.Frame(root, padx=10, pady=10)
content_frame.config(bg="white")
content_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

# Enhanced search functionality with smart summarization
def perform_smart_search():
    global current_document

    search_term = search_var.get().strip()
    if not search_term:
        messagebox.showwarning("Search Error", "Please enter a search term.")
        return

    if not current_document:
        messagebox.showwarning("Search Error", "Please load a document first by clicking 'Show Info'.")
        return

    # Gather all text from paragraphs and tables
    all_text = []
    for para in current_document.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())
    for table in current_document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(cell.text.strip())

    # Find all sentences containing the search term
    import re
    sentences = []
    for text in all_text:
        for sentence in re.split(r'(?<=[.!?])\s+', text):
            if search_term.lower() in sentence.lower():
                sentences.append(sentence.strip())

    if not sentences:
        search_summary_label.config(
            text=f"No references to '{search_term}' found in the document.",
            fg="#721c24", bg="#f8d7da"
        )
        return

    # Combine sentences into a summary of up to 100 words
    summary_words = []
    for sentence in sentences:
        for word in sentence.split():
            if len(summary_words) < 100:
                summary_words.append(word)
            else:
                break
        if len(summary_words) >= 100:
            break

    summary = " ".join(summary_words)
    if len(summary_words) == 100:
        summary += "..."

    summary = f"🔍 {summary}"

    search_summary_label.config(
        text=summary,
        fg="#0d4f1a", bg="#d4edda"
    )

    if hasattr(perform_smart_search, 'content_widgets'):
        highlight_search_matches(search_term)
    
def highlight_search_matches(search_term):
    """Highlight all occurrences of the search term in the summary and content widgets."""
    if not hasattr(perform_smart_search, 'content_widgets'):
        return
    for widget, original_text in perform_smart_search.content_widgets:
        try:
            if isinstance(widget, tk.Text):
                widget.config(state=tk.NORMAL)
                widget.tag_remove("highlight", "1.0", tk.END)
                idx = "1.0"
                while True:
                    idx = widget.search(search_term, idx, nocase=1, stopindex=tk.END)
                    if not idx:
                        break
                    lastidx = f"{idx}+{len(search_term)}c"
                    widget.tag_add("highlight", idx, lastidx)
                    idx = lastidx
                widget.tag_config("highlight", background="yellow", foreground="black")
                widget.config(state=tk.DISABLED)
            elif isinstance(widget, tk.Label):
                # For labels, just change background if search term is present
                if search_term.lower() in original_text.lower():
                    widget.config(bg="yellow", fg="black")
                else:
                    # Reset to original colors
                    widget.config(bg=widget.master.cget("bg"), fg=widget.cget("fg"))
        except Exception:
            pass

#
def create_contextual_summary(search_term, matching_paragraphs):
    """Create a smart contextual summary based on search term and found content"""
    
    if not matching_paragraphs:
        return f"❌ No content found for '{search_term}'"
    
    # Get context around search term
    relevant_content = []
    
    for para in matching_paragraphs:
        # Clean the paragraph
        clean_para = re.sub(r'\s+', ' ', para.strip())
        
        # If it's a short paragraph, include it entirely
        if len(clean_para) <= 200:
            relevant_content.append(clean_para)
        else:
            # Find sentences containing the search term
            sentences = re.split(r'[.!?]+', clean_para)
            for sentence in sentences:
                if search_term.lower() in sentence.lower() and len(sentence.strip()) > 20:
                    relevant_content.append(sentence.strip() + ".")
    
    # Remove duplicates
    unique_content = list(dict.fromkeys(relevant_content))
    
    # Special handling for "perforation" - get first 2 detailed sentences
    if search_term.lower() == "perforation":
        perforation_details = []
        for content in unique_content:
            if any(keyword in content.lower() for keyword in 
                   ['perforation', 'perforated', 'perforating', 'object-', 'testing', 'interval']):
                perforation_details.append(content)
                if len(perforation_details) >= 2:
                    break
        
        if perforation_details:
            summary_text = f"🎯 PERFORATION SUMMARY:\n\n"
            for i, detail in enumerate(perforation_details, 1):
                summary_text += f"{i}. {detail}\n\n"
            
            # Add AI summary if available - FIXED LINE 249
            if summarizer and len(" ".join(perforation_details)) > 100:
                try:
                    ai_summary = summarizer(" ".join(perforation_details), 
                                          max_length=100, min_length=30, do_sample=False)
                    summary_text += f"🤖 AI Summary: {ai_summary[0]['summary_text']}"
                except Exception as e:
                    summary_text += f"🤖 AI Summary: Could not generate summary due to error: {str(e)}"
                    
            return summary_text
    
    # General contextual summary
    if unique_content:
        summary_text = f"🔍 SEARCH RESULTS FOR '{search_term.upper()}':\n\n"
        summary_text += f"📊 Found in {len(unique_content)} contexts:\n\n"
        
        # Show top 3 most relevant results
        for i, content in enumerate(unique_content[:3], 1):
            # Highlight the search term in the context
            highlighted_content = re.sub(
                f'({re.escape(search_term)})', 
                r'📌\1📌', 
                content, 
                flags=re.IGNORECASE
            )
            summary_text += f"{i}. {highlighted_content}\n\n"
        
        if len(unique_content) > 3:
            summary_text += f"... and {len(unique_content) - 3} more references\n\n"
        
        # Add AI-generated summary if available
        if summarizer and len(" ".join(unique_content[:3])) > 100:
            try:
                combined_text = " ".join(unique_content[:3])
                ai_summary = summarizer(combined_text, max_length=80, min_length=20, do_sample=False)
                summary_text += f"🤖 AI Summary: {ai_summary[0]['summary_text']}"
            except Exception as e:
                summary_text += f"🤖 AI Summary: Could not generate summary due to error: {str(e)}"
        
        return summary_text
    
    return f"❓ Found '{search_term}' but could not extract meaningful context."

def clear_search():
    global search_results, current_search_index
    
    search_var.set("")
    search_results = []
    current_search_index = 0
    
    # Reset search summary display
    search_summary_label.config(text="Search results will appear here...", 
                               fg="#666666", bg="#f8f9fa")
    
    # Clear highlights if any
    if hasattr(perform_smart_search, 'content_widgets'):
        for widget, original_text in perform_smart_search.content_widgets:
            try:
                widget.config(bg="white" if "white" in str(widget.cget("bg")) else widget.cget("bg"))
            except:
                pass

# Enhanced perforation and testing bullet point extraction
def extract_perforation_testing_bullets(document):
    """Extract comprehensive perforation and testing information as full sentences"""
    bullet_points = []
    
    # Enhanced patterns to capture complete information
    perforation_patterns = {
        'Perforation Objects': {
            'patterns': [
                r'Object[-\s]I+V?:?\s(.+?)(?=[.;]|Object|$)',
                r'Object[-\s][1-4]:?\s(.+?)(?=[.;]|Object|$)',
                r'(?i)perforation.object.:?\s*(.+?)(?=[.;]|$)',
            ],
            'context_keywords': ['sand', 'formation', 'interval', 'depth', 'feet', 'ft']
        },
        'Testing Procedures': {
            'patterns': [
                r'(?i)formation\s+test[^.]*(.+?)(?=[.;]|$)',
                r'(?i)drill\s+stem\s+test[^.]*(.+?)(?=[.;]|$)',
                r'(?i)dst[^.]*(.+?)(?=[.;]|$)',
                r'(?i)testing\s+procedure[^.]*(.+?)(?=[.;]|$)',
            ],
            'context_keywords': ['conducted', 'performed', 'carried out', 'executed']
        },
        'Test Results': {
            'patterns': [
                r'(?i)test\s+result[^.]*(.+?)(?=[.;]|$)',
                r'(?i)testing\s+result[^.]*(.+?)(?=[.;]|$)',
                r'(?i)formation\s+response[^.]*(.+?)(?=[.;]|$)',
                r'(?i)outcome[^.]*(.+?)(?=[.;]|$)',
            ],
            'context_keywords': ['showed', 'indicated', 'revealed', 'demonstrated']
        },
        'Production Data': {
            'patterns': [
                r'(?i)production[^.](\d+.?bpd.*?)(?=[.;]|$)',
                r'(?i)flow\s+rate[^.](\d+.?)(?=[.;]|$)',
                r'(?i)gas\s+production[^.](\d+.?mcfd.*?)(?=[.;]|$)',
                r'(?i)oil\s+production[^.](\d+.?)(?=[.;]|$)',
            ],
            'context_keywords': ['barrels', 'cubic', 'per', 'day']
        },
        'Depth Information': {
            'patterns': [
                r'(?i)depth[^.](\d+.?(?:feet|ft).*?)(?=[.;]|$)',
                r'(?i)interval[^.](\d+.?(?:feet|ft).*?)(?=[.;]|$)',
                r'(?i)perforation.interval[^.](.+?)(?=[.;]|$)',
            ],
            'context_keywords': ['feet', 'ft', 'meters', 'tvd', 'md']
        },
        'Formation Details': {
            'patterns': [
                r'(?i)formation[^.]*(.+?)(?=[.;]|$)',
                r'(?i)sand[^.]formation[^.](.+?)(?=[.;]|$)',
                r'(?i)reservoir[^.]*(.+?)(?=[.;]|$)',
            ],
            'context_keywords': ['properties', 'characteristics', 'composition']
        },
        'Well Status': {
            'patterns': [
                r'(?i)well\s+status[^.]*(.+?)(?=[.;]|$)',
                r'(?i)completion\s+status[^.]*(.+?)(?=[.;]|$)',
                r'(?i)final\s+status[^.]*(.+?)(?=[.;]|$)',
                r'(?i)recommendation[^.]*(.+?)(?=[.;]|$)',
            ],
            'context_keywords': ['abandoned', 'suspended', 'completed', 'recommended']
        }
    }
    
    # Extract from paragraphs
    for para in document.paragraphs:
        text = para.text.strip()
        if len(text) < 30:  # Skip very short paragraphs
            continue
            
        text_lower = text.lower()
        
        # Check each category
        for category, info in perforation_patterns.items():
            found_match = False
            
            # Try each pattern
            for pattern in info['patterns']:
                matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
                for match in matches:
                    if found_match:
                        break
                        
                    full_sentence = create_full_sentence(text, match, category, info['context_keywords'])
                    if full_sentence and len(full_sentence) > 40:
                        bullet_points.append({
                            'category': category,
                            'text': full_sentence,
                            'source': 'paragraph',
                            'priority': get_category_priority(category)
                        })
                        found_match = True
                        break
            
            if found_match:
                continue
                
            # Fallback: Check for category keywords
            if any(keyword in text_lower for keyword in info['context_keywords']):
                # Extract sentences containing relevant keywords
                sentences = split_into_sentences(text)
                for sentence in sentences:
                    if any(keyword in sentence.lower() for keyword in info['context_keywords']):
                        full_sentence = enhance_sentence_context(sentence, category)
                        if full_sentence and len(full_sentence) > 40:
                            bullet_points.append({
                                'category': category,
                                'text': full_sentence,
                                'source': 'paragraph',
                                'priority': get_category_priority(category)
                            })
                            break
    
    # Extract from tables
    table_bullets = extract_table_full_sentences(document)
    bullet_points.extend(table_bullets)
    
    # Remove duplicates and sort
    unique_bullets = remove_duplicate_sentences(bullet_points)
    unique_bullets.sort(key=lambda x: (x['priority'], -len(x['text'])))
    
    return unique_bullets

def create_full_sentence(text, match, category, context_keywords):
    """Create a complete, contextual sentence from a regex match"""
    
    # Get the matched content
    matched_content = match.group(1) if match.groups() else match.group(0)
    
    # Find the sentence boundaries around the match
    start_pos = max(0, match.start() - 100)
    end_pos = min(len(text), match.end() + 100)
    context = text[start_pos:end_pos]
    
    # Split into sentences and find the one containing our match
    sentences = split_into_sentences(context)
    target_sentence = ""
    
    for sentence in sentences:
        if matched_content.lower() in sentence.lower():
            target_sentence = sentence
            break
    
    if not target_sentence:
        target_sentence = matched_content
    
    # Enhance the sentence with category context
    enhanced_sentence = enhance_sentence_context(target_sentence, category)
    
    return enhanced_sentence

def split_into_sentences(text):
    """Split text into sentences while preserving context"""
    # Clean the text
    text = re.sub(r'\s+', ' ', text.strip())
    
    # Split on sentence boundaries
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
    
    # Clean and filter sentences
    clean_sentences = []
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) > 15 and not sentence.startswith(('•', '-', '*')):
            # Ensure sentence ends properly
            if not sentence.endswith(('.', '!', '?')):
                sentence += '.'
            clean_sentences.append(sentence)
    
    return clean_sentences

def enhance_sentence_context(sentence, category):
    """Enhance a sentence with category-specific context"""
    
    sentence = sentence.strip()
    if not sentence:
        return ""
    
    # Category-specific enhancements
    category_prefixes = {
        'Perforation Objects': 'The perforation object',
        'Testing Procedures': 'The testing procedure',
        'Test Results': 'The test results',
        'Production Data': 'The production data',
        'Depth Information': 'The depth information',
        'Formation Details': 'The formation analysis',
        'Well Status': 'The well status'
    }
    
    # Clean the sentence
    sentence = re.sub(r'\s+', ' ', sentence)
    sentence = sentence[0].upper() + sentence[1:] if len(sentence) > 1 else sentence.upper()
    
    # Ensure proper ending
    if not sentence.endswith(('.', '!', '?')):
        sentence += '.'
    
    # Add context if the sentence is too generic
    if len(sentence) < 50 or not any(word in sentence.lower() for word in ['shows', 'indicates', 'reveals', 'demonstrates', 'contains', 'includes']):
        prefix = category_prefixes.get(category, 'The analysis')
        if not sentence.lower().startswith(('the ', 'this ', 'it ', 'that ')):
            sentence = f"{prefix} shows that {sentence.lower()}"
        else:
            sentence = f"{prefix} indicates {sentence.lower()}"
    
    return sentence

def extract_table_full_sentences(document):
    """Extract full sentences from tables"""
    table_bullets = []
    
    for table_idx, table in enumerate(document.tables):
        try:
            if not table.rows:
                continue
                
            headers = []
            # Get headers
            if table.rows:
                for cell in table.rows[0].cells:
                    header_text = cell.text.strip()
                    if header_text:
                        headers.append(header_text)
            
            # Process data rows
            for row_idx, row in enumerate(table.rows[1:], 1):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                
                if row_data and len(row_data) >= 2:
                    # Create structured sentence
                    sentence = create_table_sentence(headers, row_data, table_idx, row_idx)
                    if sentence:
                        table_bullets.append({
                            'category': 'Tabulated Data',
                            'text': sentence,
                            'source': f'table_{table_idx + 1}',
                            'priority': 9
                        })
        
        except Exception as e:
            print(f"Error processing table {table_idx}: {e}")
            continue
    
    return table_bullets

def create_table_sentence(headers, row_data, table_idx, row_idx):
    """Create a full sentence from table row data"""
    
    if not headers or not row_data:
        return ""
    
    # Create structured information
    structured_parts = []
    for i, data in enumerate(row_data):
        if i < len(headers) and data.strip():
            header = headers[i].strip()
            value = data.strip()
            
            # Create natural language parts
            if 'object' in header.lower():
                structured_parts.append(f"the {header.lower()} is {value}")
            elif 'interval' in header.lower() or 'depth' in header.lower():
                structured_parts.append(f"the {header.lower()} spans {value}")
            elif 'result' in header.lower() or 'outcome' in header.lower():
                structured_parts.append(f"the {header.lower()} shows {value}")
            elif 'test' in header.lower():
                structured_parts.append(f"the {header.lower()} indicates {value}")
            else:
                structured_parts.append(f"the {header.lower()} is {value}")
    
    if structured_parts:
        # Join parts into a coherent sentence
        if len(structured_parts) == 1:
            sentence = f"The table data shows that {structured_parts[0]}."
        elif len(structured_parts) == 2:
            sentence = f"The analysis reveals that {structured_parts[0]} and {structured_parts[1]}."
        else:
            main_parts = structured_parts[:2]
            additional = len(structured_parts) - 2
            sentence = f"The comprehensive data indicates that {', '.join(main_parts)}, along with {additional} additional parameter{'s' if additional > 1 else ''}."
        
        # Capitalize first letter and ensure proper ending
        sentence = sentence[0].upper() + sentence[1:] if len(sentence) > 1 else sentence.upper()
        if not sentence.endswith('.'):
            sentence += '.'
            
        return sentence
    
    return ""

def get_category_priority(category):
    """Get priority for category sorting"""
    priorities = {
        'Perforation Objects': 1,
        'Testing Procedures': 2,
        'Test Results': 3,
        'Production Data': 4,
        'Depth Information': 5,
        'Formation Details': 6,
        'Well Status': 7,
        'Tabulated Data': 8
    }
    return priorities.get(category, 9)

def remove_duplicate_sentences(bullet_points):
    """Remove duplicate sentences while preserving the best versions"""
    unique_bullets = []
    seen_content = set()
    
    for bullet in bullet_points:
        # Normalize for comparison
        normalized = re.sub(r'[^\w\s]', '', bullet['text'].lower())
        normalized = re.sub(r'\s+', ' ', normalized).strip()
        
        # Create a shorter key for similarity checking
        key_words = normalized.split()[:8]  # First 8 words
        content_key = ' '.join(key_words)
        
        if len(content_key) > 20 and content_key not in seen_content:
            seen_content.add(content_key)
            unique_bullets.append(bullet)
    
    return unique_bullets

def create_full_sentence_summary(bullet_points):
    """Create a comprehensive summary with full sentences in bullet points"""
    
    if not bullet_points:
        return "• No specific perforation or testing information found in the document."
    
    # Group by category
    categorized_bullets = {}
    for bullet in bullet_points:
        category = bullet['category']
        if category not in categorized_bullets:
            categorized_bullets[category] = []
        categorized_bullets[category].append(bullet['text'])
    
    # Build the summary
    summary_parts = []
    
    # Header
    total_sentences = len(bullet_points)
    categories_count = len(categorized_bullets)
    
    summary_parts.append("📋 COMPREHENSIVE PERFORATION & TESTING ANALYSIS")
    summary_parts.append("═" * 55)
    summary_parts.append(f"📊 Summary: {total_sentences} detailed findings across {categories_count} technical categories")
    summary_parts.append("")
    
    # Category icons
    category_icons = {
        'Perforation Objects': '🎯',
        'Testing Procedures': '🔬',
        'Test Results': '📈',
        'Production Data': '⚡',
        'Depth Information': '📏',
        'Formation Details': '🗻',
        'Well Status': '📋',
        'Tabulated Data': '📊'
    }
    
    # Present each category with full sentences
    for category, sentences in categorized_bullets.items():
        if sentences:
            icon = category_icons.get(category, '•')
            summary_parts.append(f"{icon} {category.upper()}:")
            summary_parts.append("─" * (len(category) + 5))
            
            # Show up to 5 sentences per category
            for i, sentence in enumerate(sentences[:5], 1):
                # Ensure it's a complete sentence
                clean_sentence = sentence.strip()
                if not clean_sentence.endswith('.'):
                    clean_sentence += '.'
                
                summary_parts.append(f"   • {clean_sentence}")
            
            if len(sentences) > 5:
                summary_parts.append(f"   • ... and {len(sentences) - 5} additional detailed findings in this category.")
            
            summary_parts.append("")  # Add space between categories
    
    # Technical summary
    summary_parts.append("🔍 TECHNICAL SUMMARY:")
    summary_parts.append("═" * 25)
    
    # Count different types of information
    tech_counts = {
        'Object References': len([b for b in bullet_points if 'object' in b['text'].lower()]),
        'Test Procedures': len([b for b in bullet_points if any(word in b['text'].lower() for word in ['test', 'testing', 'procedure'])]),
        'Numerical Data': len([b for b in bullet_points if re.search(r'\d+', b['text'])]),
        'Depth Measurements': len([b for b in bullet_points if any(word in b['text'].lower() for word in ['feet', 'ft', 'depth', 'interval'])]),
        'Production Metrics': len([b for b in bullet_points if any(word in b['text'].lower() for word in ['bpd', 'mcfd', 'production', 'flow'])])
    }
    
    for metric, count in tech_counts.items():
        if count > 0:
            summary_parts.append(f"   • {metric}: {count} comprehensive statements identified")
    
    # Key insights
    summary_parts.append("")
    summary_parts.append("⭐ KEY INSIGHTS:")
    summary_parts.append("═" * 18)
    
    # Extract most informative sentences
    key_insights = []
    for bullet in bullet_points[:3]:  # Top 3 by priority
        sentence = bullet['text'].strip()
        if len(sentence) > 60:  # Substantial content
            key_insights.append(f"   • {sentence}")
    
    if key_insights:
        summary_parts.extend(key_insights)
    else:
        summary_parts.append("   • All findings provide comprehensive technical details about the perforation and testing operations.")
    
    return "\n".join(summary_parts)

# Modified Show Info function with fixed table display
def show_perforation_info():
    global current_document, document_paragraphs
    
    field = field_var.get()
    well = well_var.get()
    keyword = keyword_var.get()
    
    if not field or not well:
        messagebox.showwarning("Input Error", "Please select both field and well name.")
        return

    # Clear right panel
    for w in content_frame.winfo_children():
        w.destroy()

    # Find the matching document path
    path = None
    for wname, full_path in field_well_map.get(field, []):
        if wname == well:
            path = full_path
            break
    
    if not path:
        messagebox.showerror("Not Found", "Matching document not found.")
        return

    try:
        document = Document(path)
        current_document = document
        
        # Store all paragraphs for search functionality
        document_paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        
        # Create header with document info
        header_frame = tk.Frame(content_frame, bg="#f0f0f0", relief="solid", bd=1)
        header_frame.pack(fill="x", padx=10, pady=(10,0))
        
        doc_info_frame = tk.Frame(header_frame, bg="#f0f0f0")
        doc_info_frame.pack(side="left", fill="x", expand=True, padx=10, pady=5)
        
        doc_title = tk.Label(doc_info_frame, text=f"🔍 Comprehensive Analysis for '{keyword}'", 
                            font=("Segoe UI", 12, "bold"), bg="#f0f0f0")
        doc_title.pack(anchor="w")
        
        doc_details = tk.Label(doc_info_frame, text=f"Field: {field} | Well: {well} | Document: {os.path.basename(path)}", 
                              font=("Segoe UI", 10), bg="#f0f0f0", fg="gray")
        doc_details.pack(anchor="w")
        
        main_frame = tk.Frame(content_frame)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        canvas = tk.Canvas(main_frame, bg="white")
        scrollbar = tk.Scrollbar(sidebar, orient=tk.VERTICAL)
        scrollable_frame = tk.Frame(canvas, bg="white")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        def configure_canvas_width(event=None):
            canvas_width = event.width if event else canvas.winfo_width()
            canvas.itemconfig(canvas.create_window((0, 0), window=scrollable_frame, anchor="nw"), width=canvas_width)

        canvas.bind('<Configure>', configure_canvas_width)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        perform_smart_search.content_widgets = []
        
        loading_frame = tk.Frame(scrollable_frame, bg="#fff3cd", relief="solid", bd=1)
        loading_frame.pack(fill="x", padx=20, pady=10)
        
        loading_label = tk.Label(loading_frame, text="🤖 Creating Full Sentence Analysis...", 
                               font=("Calibri", 12, "bold"), bg="#fff3cd", fg="#856404")
        loading_label.pack(pady=10)
        
        root.update()
        
        bullet_points = extract_perforation_testing_bullets(document)
        
        if bullet_points:
            full_sentence_summary = create_full_sentence_summary(bullet_points)
            
            loading_frame.destroy()
            
            summary_frame = tk.Frame(scrollable_frame, bg="#e8f5e8", relief="solid", bd=2)
            summary_frame.pack(fill="x", padx=20, pady=10)
            
            summary_title = tk.Label(summary_frame, text="📋 PERFORATION & TESTING - FULL SENTENCE ANALYSIS", 
                                   font=("Calibri", 16, "bold"), bg="#e8f5e8", fg="#155724")
            summary_title.pack(anchor="w", padx=15, pady=(10, 5))
            
            summary_subtitle = tk.Label(summary_frame, text="🎯 Complete Sentences Providing Comprehensive Technical Details", 
                                      font=("Calibri", 11, "italic"), bg="#e8f5e8", fg="#155724")
            summary_subtitle.pack(anchor="w", padx=15, pady=(0, 10))
            
            text_frame = tk.Frame(summary_frame, bg="#e8f5e8")
            text_frame.pack(fill="both", expand=True, padx=15, pady=(0, 10))
            
            summary_text = tk.Text(text_frame, 
                                 font=("Calibri", 11), 
                                 bg="#f8fff8", 
                                 fg="#155724",
                                 wrap=tk.WORD, 
                                 height=25,
                                 relief="solid",
                                 bd=1,
                                 padx=15,
                                 pady=15)
            
            text_scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=summary_text.yview)
            summary_text.configure(yscrollcommand=text_scrollbar.set)
            
            summary_text.pack(side="left", fill="both", expand=True)
            text_scrollbar.pack(side="right", fill="y")
            
            summary_text.insert(tk.END, full_sentence_summary)
            summary_text.config(state=tk.DISABLED)
            
            perform_smart_search.content_widgets.append((summary_title, summary_title.cget("text")))
            perform_smart_search.content_widgets.append((summary_text, full_sentence_summary))
            
            stats_frame = tk.Frame(summary_frame, bg="#d4edda", relief="solid", bd=1)
            stats_frame.pack(fill="x", padx=15, pady=(0, 10))
            
            stats_text = f"📊 Full Sentence Analysis: {len(bullet_points)} complete statements | Categories: {len(set(bp['category'] for bp in bullet_points))} | Format: Complete sentences with full context"
            
            stats_label = tk.Label(stats_frame, text=stats_text, 
                                 font=("Calibri", 9, "bold"), bg="#d4edda", fg="#0d4f1a")
            stats_label.pack(anchor="w", padx=10, pady=8)
        else:
            loading_frame.destroy()
            
            no_perf_frame = tk.Frame(scrollable_frame, bg="#f8d7da", relief="solid", bd=1)
            no_perf_frame.pack(fill="x", padx=20, pady=10)
            
            no_perf_label = tk.Label(no_perf_frame, text="⚠ No specific perforation or testing content found for detailed analysis", 
                                   font=("Calibri", 11, "bold"), bg="#f8d7da", fg="#721c24")
            no_perf_label.pack(pady=10)
        
        
        found_content = False
        
        for i, table in enumerate(document.tables):
            table_contains_objects = False
            table_text = ""
            
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.lower()
                    table_text += cell_text + " "
                    if ("object-i" in cell_text or "object-ii" in cell_text or 
                        "object-iii" in cell_text or "object-iv" in cell_text or
                        ("object" in cell_text and "interval" in table_text) or
                        ("sand" in cell_text and "testing result" in table_text)):
                        table_contains_objects = True
                        break
                if table_contains_objects:
                    break
            
            if table_contains_objects:
                table_widgets = insert_professional_table(scrollable_frame, table, i+1)
                if table_widgets:
                    perform_smart_search.content_widgets.extend(table_widgets)
                found_content = True
                break
        
        if not found_content:
            for i, table in enumerate(document.tables):
                table_has_testing = False
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.lower()
                        if ("testing" in cell_text or "interval" in cell_text or 
                            "result" in cell_text or "formation" in cell_text):
                            table_has_testing = True
                            break
                    if table_has_testing:
                        break
                
                if table_has_testing:
                    table_widgets = insert_professional_table(scrollable_frame, table, i+1)
                    if table_widgets:
                        perform_smart_search.content_widgets.extend(table_widgets)
                    found_content = True
                    break
        
        if not found_content:
            no_content_frame = tk.Frame(scrollable_frame, bg="white")
            no_content_frame.pack(fill="x", padx=20, pady=20)
            
            no_content_label = tk.Label(no_content_frame, text="📋 No perforation testing tables found in this document.", 
                                       font=("Calibri", 12), bg="white", fg="#666666")
            no_content_label.pack(anchor="center")
            perform_smart_search.content_widgets.append((no_content_label, no_content_label.cget("text")))
        
    except Exception as e:
        messagebox.showerror("Error", f"Failed to read document.\nError: {str(e)}")
        print(f"Detailed error: {e}")

def insert_professional_table(parent_frame, table, table_number):
    """Insert a professional, properly formatted table"""
    search_widgets = []
    
    try:
        if not table.rows:
            return search_widgets

        table_data = []
        processed_cells = set()
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_pos = f"{row_idx}_{cell_idx}"
                
                if cell_pos in processed_cells:
                    continue
                
                cell_text = ""
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_text += paragraph.text.strip() + " "
                
                cell_text = cell_text.strip()
                row_data.append(cell_text)
                processed_cells.add(cell_pos)
            
            if any(cell.strip() for cell in row_data):
                table_data.append(row_data)

        if not table_data:
            return search_widgets

        filtered_data = []
        for row_data in table_data:
            row_text = " ".join(row_data).lower()
            if any(skip_phrase in row_text for skip_phrase in [
                "operation with drilling rig", "objective: location", "prospect of surma",
                "drilled down", "casing was lowered", "mud was further increased",
                "to explore the hydrocarbon", "rig: e-1400-x"
            ]):
                continue
            
            max_cols = max(len(row) for row in table_data) if table_data else 4
            while len(row_data) < max_cols:
                row_data.append("")
            
            filtered_data.append(row_data)

        if not filtered_data:
            return search_widgets

        table_container = tk.Frame(parent_frame, bg="white", relief="solid", bd=2)
        table_container.pack(fill="x", padx=20, pady=10)

        title_frame = tk.Frame(table_container, bg="#4472C4")
        title_frame.pack(fill="x")
        
        table_title = tk.Label(title_frame, 
                              text=f"📊 TABLE {table_number}: PERFORATION & TESTING DATA", 
                              font=("Calibri", 14, "bold"), 
                              bg="#4472C4", fg="white",
                              pady=10)
        table_title.pack()
        search_widgets.append((table_title, table_title.cget("text")))

        table_frame = tk.Frame(table_container, bg="white")
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        canvas = tk.Canvas(table_frame, bg="white", height=400)
        h_scrollbar = ttk.Scrollbar(table_frame, orient="horizontal", command=canvas.xview)
        v_scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=canvas.yview)
        
        canvas.configure(xscrollcommand=h_scrollbar.set, yscrollcommand=v_scrollbar.set)

        content_frame_table = tk.Frame(canvas, bg="white")
        
        max_cols = len(filtered_data[0]) if filtered_data else 4
        col_widths = []
        
        for col_idx in range(max_cols):
            max_width = 100  
            for row_data in filtered_data:
                if col_idx < len(row_data):
                    text_width = len(str(row_data[col_idx])) * 8  
                    max_width = max(max_width, min(text_width, 300))  
            col_widths.append(max_width)

        for row_idx, row_data in enumerate(filtered_data):
            for col_idx in range(max_cols):
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                
                if row_idx == 0:  
                    bg_color = "#4472C4"
                    fg_color = "white"
                    font_style = ("Arial", 10, "bold")
                    relief_style = "solid"
                    bd_width = 2
                elif "Object-" in cell_text:
                    bg_color = "#E7F3FF"
                    fg_color = "#0066CC"
                    font_style = ("Arial", 10, "bold")
                    relief_style = "solid"
                    bd_width = 1
                else:  
                    bg_color = "#FAFAFA" if row_idx % 2 == 0 else "white"
                    fg_color = "black"
                    font_style = ("Arial", 9, "normal")
                    relief_style = "solid"
                    bd_width = 1
                
                cell_widget = tk.Label(
                    content_frame_table,
                    text=cell_text,
                    font=font_style,
                    bg=bg_color,
                    fg=fg_color,
                    relief=relief_style,
                    bd=bd_width,
                    anchor="w" if col_idx > 0 else "center",
                    justify="left",
                    padx=8,
                    pady=6,
                    wraplength=col_widths[col_idx] - 20
                )
                
                cell_widget.grid(row=row_idx, column=col_idx, sticky="nsew", padx=0, pady=0)
                
                content_frame_table.grid_columnconfigure(col_idx, weight=0, minsize=col_widths[col_idx])
                
                if cell_text.strip():
                    search_widgets.append((cell_widget, cell_text))

        canvas.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

        canvas_window = canvas.create_window((0, 0), window=content_frame_table, anchor="nw")

        def configure_scroll_region(event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
            req_height = content_frame_table.winfo_reqheight()
            canvas_height = min(req_height + 20, 400)  # Max height 400px
            canvas.configure(height=canvas_height)

        def configure_canvas_window(event=None):
            canvas_width = canvas.winfo_width()
            req_width = content_frame_table.winfo_reqwidth()
            if req_width <= canvas_width:
                canvas.itemconfig(canvas_window, width=canvas_width)

        content_frame_table.bind("<Configure>", configure_scroll_region)
        canvas.bind("<Configure>", configure_canvas_window)

        def on_mouse_wheel(event):
            if event.state & 0x1:  # Shift key pressed
                canvas.xview_scroll(int(-1 * (event.delta / 120)), "units")
            else:
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            return "break"

        def bind_mouse_wheel(widget):
            widget.bind("<MouseWheel>", on_mouse_wheel)
            for child in widget.winfo_children():
                bind_mouse_wheel(child)

        bind_mouse_wheel(canvas)
        bind_mouse_wheel(content_frame_table)

        canvas.bind("<Button-1>", lambda e: canvas.focus_set())
        canvas.bind("<Key>", lambda e: handle_key_navigation(e, canvas))

        table_container.update_idletasks()
        configure_scroll_region()

        summary_frame = tk.Frame(table_container, bg="#F0F8FF", relief="solid", bd=1)
        summary_frame.pack(fill="x", padx=5, pady=(0, 5))
        
        rows_count = len(filtered_data) - 1  
        cols_count = max_cols
        
        summary_text = f"📋 Table Summary: {rows_count} data rows × {cols_count} columns | Perforation testing details with comprehensive technical specifications"
        
        summary_label = tk.Label(summary_frame, text=summary_text,
                                font=("Arial", 9, "italic"), bg="#F0F8FF", fg="#0066CC",
                                padx=10, pady=5, wraplength=800, justify="left")
        summary_label.pack(anchor="w")
        search_widgets.append((summary_label, summary_text))

        return search_widgets

    except Exception as e:
        print(f"Professional table error: {e}")
        
        # Fallback: Create error message
        error_frame = tk.Frame(parent_frame, bg="#FFE6E6", relief="solid", bd=1)
        error_frame.pack(fill="x", padx=20, pady=10)
        
        error_label = tk.Label(error_frame, text=f"⚠ Error displaying table {table_number}: {str(e)[:100]}...",
                              font=("Arial", 10), bg="#FFE6E6", fg="#CC0000",
                              padx=10, pady=10, wraplength=800, justify="left")
        error_label.pack()
        
        return [(error_label, error_label.cget("text"))]

def handle_key_navigation(event, canvas):
    """Handle keyboard navigation for table scrolling"""
    if event.keysym == "Up":
        canvas.yview_scroll(-1, "units")
    elif event.keysym == "Down":
        canvas.yview_scroll(1, "units")
    elif event.keysym == "Left":
        canvas.xview_scroll(-1, "units")
    elif event.keysym == "Right":
        canvas.xview_scroll(1, "units")
    elif event.keysym == "Prior":  # Page Up
        canvas.yview_scroll(-1, "pages")
    elif event.keysym == "Next":   # Page Down
        canvas.yview_scroll(1, "pages")
    elif event.keysym == "Home":
        canvas.yview_moveto(0)
        canvas.xview_moveto(0)
    elif event.keysym == "End":
        canvas.yview_moveto(1)

# Other functions remain the same...
def open_doc_externally():
    field = field_var.get()
    well = well_var.get()
    
    if not field or not well:
        messagebox.showwarning("Input Error", "Please select both field and well name.")
        return

    path = None
    for wname, full_path in field_well_map.get(field, []):
        if wname == well:
            path = full_path
            break
    
    if not path:
        messagebox.showerror("Not Found", "Matching document not found.")
        return

    try:
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":
            subprocess.run(["open", path])
        else:
            subprocess.run(["xdg-open", path])
            
        messagebox.showinfo("Success", f"Document opened: {os.path.basename(path)}")
        
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open document.\nError: {str(e)}")

def open_doc_folder():
    field = field_var.get()
    well = well_var.get()
    
    if not field or not well:
        messagebox.showwarning("Input Error", "Please select both field and well name.")
        return

    path = None
    for wname, full_path in field_well_map.get(field, []):
        if wname == well:
            path = full_path
            break
    
    if not path:
        messagebox.showerror("Not Found", "Matching document not found.")
        return

    try:
        folder_path = os.path.dirname(path)
        if platform.system() == "Windows":
            subprocess.run(["explorer", folder_path])
        elif platform.system() == "Darwin":
            subprocess.run(["open", folder_path])
        else:
            subprocess.run(["xdg-open", folder_path])
            
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open folder.\nError: {str(e)}")

def clear_content():
    global current_document, document_paragraphs
    
    for w in content_frame.winfo_children():
        w.destroy()
    well_dropdown.set("")
    field_dropdown.set("")
    keyword_dropdown.set("Perforation")
    
    current_document = None
    document_paragraphs = []
    clear_search()

def on_field_select(event):
    field = field_var.get()
    if field in field_well_map:
        wells = [well_name for well_name, _ in field_well_map[field]]
        well_dropdown.config(values=wells)
        well_dropdown.set("")

# Bind events
field_dropdown.bind("<<ComboboxSelected>>", on_field_select)
search_entry.bind("<Return>", lambda event: perform_smart_search())

# “Search & Summarize” button: ~18 chars wide
search_btn = tk.Button(
    search_btn_frame,
    text="🔍 Search & Summarize",
    command= perform_smart_search,
    bg="#28a745", fg="white",
    font=("Segoe UI", 9),
    width=18        
)
search_btn.pack(side="left", padx=(0, 4))   


# “Clear” button: ~8 chars wide
clear_search_btn = tk.Button(
    search_btn_frame,
    text="Clear",
    command=clear_content,
    bg="#6c757d", fg="white",
    font=("Segoe UI", 9),
    width=8                                   
)
clear_search_btn.pack(side="left")        

search_summary_frame = tk.Frame(
    controls_frame,
    bg="#f8f9fa", relief="solid", bd=1,
    width=280             
)
search_summary_frame.pack(fill="x", padx=10, pady=5)

search_summary_label = tk.Label(
    search_summary_frame,
    text="Search results will appear here...",
    font=("Segoe UI", 9),
    bg="#f8f9fa", fg="#666666",
    wraplength=260,                  
    justify="left"
)
search_summary_label.pack(padx=10, pady=8)


# Show Info (~18 chars wide)
show_table_btn = tk.Button(
    controls_frame,
    text="🔍 Analyze Report",
    command=show_perforation_info,
    bg="#007bff", fg="white",
    font=("Segoe UI", 10, "bold"),
    width=18                              
)
show_table_btn.pack(fill="x", padx=10, pady=8)
# Download PDF (~14 chars wide)
download_pdf_btn = tk.Button(
    controls_frame,
    text="📥 Download PDF",
    command=export_summary_to_pdf,
    bg="#10b981", fg="white",
    font=("Segoe UI", 10, "bold"),
    width=16
)
download_pdf_btn.pack(fill="x", padx=10, pady=5)

# Clear/Reset (~14 chars wide)
clear_btn = tk.Button(
    controls_frame,
    text="🔄 Reset Selection",
    command=clear_content,
    bg="#4b9ede", fg="white",
    width=16                             
)
clear_btn.pack(fill="x", padx=10, pady=5)

# Open in Word (~16 chars wide)
open_word_btn = tk.Button(
    controls_frame,
    text="📄 Open Report in Word",
    command=open_doc_externally,
    bg="#0078d4", fg="white",
    font=("Segoe UI", 10, "bold"),
    width=16                              
)
open_word_btn.pack(fill="x", padx=10, pady=5)

# Open Folder (~12 chars wide)
open_folder_btn = tk.Button(
    controls_frame,
    text="📁 Open Report Folder",
    command=open_doc_folder,
    bg="#6a6a6a", fg="white",
    font=("Segoe UI", 10),
    width=12                           
)
open_folder_btn.pack(fill="x", padx=10, pady=5)

# Exit (~8 chars wide)
exit_btn = tk.Button(
    controls_frame,
    text="Exit Application",
    command=root.quit,
    bg="#d32f2f", fg="white",
    font=("Segoe UI", 10, "bold"),
    width=8                            
)
exit_btn.pack(fill="x", padx=10, pady=5)

# Start the GUI
root.mainloop()